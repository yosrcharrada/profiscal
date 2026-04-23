"""
create_reports.py
Generate realistic Tunisian tax consultation reports grounded in Neo4j GraphRAG chunks.

Pipeline per report
-------------------
1. Pick one of 10 predefined, distinct consultation scenarios.
2. Query Neo4j for the most relevant legal chunks (keyword + topic scoring).
3. Build a GraphRAG context string from those chunks (real law text).
4. Call Azure OpenAI with:
   - system prompt: expert fiscal tunisien senior
   - user prompt:   scenario brief + real legal extracts + section instructions
5. Parse LLM response into template tokens.
6. Fill template_fr.docx (token replacement) → output .docx.
   Fallback: create a bare .docx from scratch if template is missing.

Environment variables (same as .NET GraphRagService):
  OPENAI_API_KEY, OPENAI_ENDPOINT, OPENAI_API_VERSION (default 2024-02-15-preview)
  LLM_MODEL (default gpt-4o)
  NEO4J_URI, NEO4J_DATABASE, NEO4J_USERNAME, NEO4J_PASSWORD
"""

from __future__ import annotations

import argparse
import hashlib
import os
import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

try:
    from dotenv import load_dotenv
except ImportError:  # pragma: no cover
    def load_dotenv() -> None:
        return None

try:
    from neo4j import GraphDatabase
except ImportError:  # pragma: no cover
    GraphDatabase = None  # type: ignore

from config import get_config

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:  # pragma: no cover
    Document = None  # type: ignore
    Pt = None  # type: ignore
    RGBColor = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore

try:
    from openai import AzureOpenAI
except ImportError:  # pragma: no cover
    AzureOpenAI = None  # type: ignore

load_dotenv()

# Max characters kept per Neo4j chunk when building the LLM prompt.
# Balances context richness with prompt token budget (~150 tokens per chunk).
_MAX_CHUNK_TEXT_CHARS = 800

# Hash prefix length used for deduplication of retrieved chunks.
# 16 hex chars = 64-bit collision space, sufficient for a small result set.
_DEDUP_HASH_LENGTH = 16


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class LegalSource:
    doc_name: str
    page_num: int
    article_ref: str
    section_title: str
    text: str
    score: float

    def short_ref(self) -> str:
        article = self.article_ref.strip() if self.article_ref else ""
        parts = [self.doc_name, f"p.{self.page_num}"]
        if article:
            parts.append(article)
        return " — ".join(parts)


@dataclass
class ConsultationScenario:
    """All the facts needed to generate one consultation."""
    idx: int
    reference: str
    case_label: str           # short tag for filenames
    client_type: str          # "entreprise" | "particulier" | "mixte"
    client_name: str
    client_description: str   # one paragraph describing the client
    situation: str            # the concrete fact pattern (free text)
    objective: str            # what the client wants to understand
    topics: Tuple[str, ...]   # Neo4j keyword topics
    extra_topics: Tuple[str, ...] = field(default_factory=tuple)  # secondary query


# ---------------------------------------------------------------------------
# 10 distinct scenarios
# ---------------------------------------------------------------------------

SCENARIOS: Tuple[ConsultationScenario, ...] = (
    ConsultationScenario(
        idx=1,
        reference="CONS-2024-001",
        case_label="saas_usa",
        client_type="entreprise",
        client_name="TechPark Innovations SARL",
        client_description=(
            "TechPark Innovations SARL est une société tunisienne spécialisée dans le "
            "développement de solutions ERP pour PME, immatriculée à Tunis, soumise à "
            "l'impôt sur les sociétés au taux normal."
        ),
        situation=(
            "La société souhaite acquérir un abonnement annuel à un logiciel SaaS "
            "(Microsoft Azure & Salesforce) auprès de fournisseurs résidant aux États-Unis, "
            "pour un montant estimé à 480 000 TND/an. Aucun établissement stable de ces "
            "fournisseurs n'existe en Tunisie."
        ),
        objective=(
            "Déterminer le traitement TVA applicable (territorialité, auto-liquidation), "
            "l'obligation de retenue à la source sur les paiements à l'étranger, et les "
            "obligations déclaratives associées (CDPF, formulaires douaniers)."
        ),
        topics=("tva", "retenue", "source", "non-résident", "prestation", "service", "états-unis"),
        extra_topics=("importation", "cdpf", "ctva", "territoire"),
    ),
    ConsultationScenario(
        idx=2,
        reference="CONS-2024-002",
        case_label="equipment_germany",
        client_type="entreprise",
        client_name="Carthage Industries SA",
        client_description=(
            "Carthage Industries SA est un groupe industriel tunisien fabriquant des "
            "composants mécaniques pour l'automobile, coté à la Bourse des Valeurs "
            "Mobilières de Tunis, réalisant un chiffre d'affaires annuel de 45 M TND."
        ),
        situation=(
            "La société prévoit l'acquisition de machines-outils CNC auprès d'un "
            "fabricant allemand (Trumpf GmbH) pour 2,1 M EUR. Le contrat prévoit en outre "
            "une prestation de mise en service et de formation réalisée par des techniciens "
            "allemands sur site tunisien pendant 6 semaines."
        ),
        objective=(
            "Analyser le régime TVA à l'importation, la retenue à la source sur la "
            "prestation de services, les droits de douane, et l'incidence éventuelle de "
            "la convention fiscale franco-allemande/tunisienne."
        ),
        topics=("tva", "importation", "retenue", "source", "douane", "allemagne", "prestation"),
        extra_topics=("convention", "établissement stable", "cdpf"),
    ),
    ConsultationScenario(
        idx=3,
        reference="CONS-2024-003",
        case_label="dg_francais",
        client_type="particulier",
        client_name="M. Laurent Dupont",
        client_description=(
            "M. Laurent Dupont est un ressortissant français, résidant à Paris, "
            "nommé Directeur Général d'une filiale tunisienne d'un groupe pharmaceutique "
            "français (Pharma Med Tunisie SARL). Il intervient en Tunisie 3 jours par "
            "semaine et perçoit sa rémunération exclusivement de la maison mère en France."
        ),
        situation=(
            "M. Dupont perçoit un salaire mensuel de 18 000 EUR versé par la maison mère "
            "française. La filiale tunisienne lui rembourse uniquement ses frais de déplacement. "
            "Il n'a pas de visa de travail tunisien ni de carte de séjour."
        ),
        objective=(
            "Déterminer son statut fiscal en Tunisie (résident / non-résident), "
            "l'imposition de sa rémunération au titre de l'IRPP tunisien, "
            "le régime des retenues à la source de la filiale, "
            "son statut au regard de la sécurité sociale, du change, "
            "et la nécessité d'un visa de travail."
        ),
        topics=("irpp", "résidence", "retenue", "source", "convention", "france", "sécurité sociale"),
        extra_topics=("directeur", "non-résident", "change", "visa", "travail"),
    ),
    ConsultationScenario(
        idx=4,
        reference="CONS-2024-004",
        case_label="remote_worker_uae",
        client_type="particulier",
        client_name="Mme Sana Mejri",
        client_description=(
            "Mme Sana Mejri est une développeuse tunisienne résidant à Tunis, "
            "travaillant en tant que consultante indépendante pour une entreprise "
            "tech basée à Dubaï (UAE). Elle facture directement en USD depuis la Tunisie."
        ),
        situation=(
            "Mme Mejri réside en Tunisie plus de 183 jours par an. Elle perçoit 5 000 USD "
            "par mois d'une société émiratie sans présence en Tunisie. Les fonds sont "
            "rapatriés via virement bancaire international sur un compte tunisien. "
            "Elle n'a ni statut d'entreprise, ni inscription à la patente."
        ),
        objective=(
            "Établir son obligation déclarative IRPP en Tunisie, le régime de change "
            "applicable (obligation de rapatriement), la TVA sur prestations exportées, "
            "et l'opportunité de créer une structure juridique (SARL, entreprise individuelle)."
        ),
        topics=("irpp", "résidence", "change", "tva", "exportation", "prestation", "indépendant"),
        extra_topics=("émirats", "convention", "rapatriement", "devises"),
    ),
    ConsultationScenario(
        idx=5,
        reference="CONS-2024-005",
        case_label="management_fees",
        client_type="entreprise",
        client_name="Groupe Méditerranée Holding (filiale tunisienne : Med Services SARL)",
        client_description=(
            "Med Services SARL est la filiale tunisienne d'un groupe français coté. "
            "La société mère française facture des « management fees » (services de "
            "direction générale, IT, finance) à la filiale tunisienne depuis 3 ans."
        ),
        situation=(
            "La maison mère française (Méditerranée Holding SA) facture annuellement "
            "1,2 M TND de management fees à Med Services SARL, soit 8 % du chiffre "
            "d'affaires de la filiale. Aucune documentation de prix de transfert n'a "
            "été établie à ce jour. L'administration fiscale a ouvert un contrôle."
        ),
        objective=(
            "Analyser la déductibilité des management fees au titre de l'IS, "
            "les exigences documentaires prix de transfert (art. CDPF), "
            "le risque de redressement, la retenue à la source applicable, "
            "et les mesures correctives à mettre en place."
        ),
        topics=("prix de transfert", "management fees", "déductibilité", "is", "documentation"),
        extra_topics=("intra-groupe", "comparables", "retenue", "cdpf", "contrôle"),
    ),
    ConsultationScenario(
        idx=6,
        reference="CONS-2024-006",
        case_label="ip_royalties",
        client_type="entreprise",
        client_name="TunisTech SARL",
        client_description=(
            "TunisTech SARL est une PME tunisienne éditrice de logiciels, filiale à 70 % "
            "d'une holding néerlandaise (NLTech BV). La holding détient la propriété "
            "intellectuelle (marque, logiciels) et autorise la filiale à les exploiter."
        ),
        situation=(
            "TunisTech SARL verse chaque année 15 % de son chiffre d'affaires à NLTech BV "
            "au titre de redevances de licence de marque et de logiciels. Le taux a été "
            "fixé par contrat sans étude de comparables. NLTech BV n'a pas d'établissement "
            "stable en Tunisie."
        ),
        objective=(
            "Apprécier la déductibilité des redevances (IS), la retenue à la source "
            "sur les paiements (convention Tunisie-Pays-Bas), le risque prix de transfert "
            "sur le taux de 15 %, la TVA applicable, et les obligations documentaires."
        ),
        topics=("redevance", "propriété intellectuelle", "retenue", "source", "convention", "pays-bas"),
        extra_topics=("prix de transfert", "is", "tva", "déductibilité"),
    ),
    ConsultationScenario(
        idx=7,
        reference="CONS-2024-007",
        case_label="doctor_sarl",
        client_type="mixte",
        client_name="Dr. Hedi Ben Salah / Clinique Lumière SARL",
        client_description=(
            "Dr. Hedi Ben Salah est médecin spécialiste (cardiologue) exerçant à la fois "
            "en cabinet libéral à titre personnel et en tant que gérant associé unique "
            "de Clinique Lumière SARL, clinique privée qu'il a créée en 2018."
        ),
        situation=(
            "Dr. Ben Salah perçoit des honoraires personnels pour consultations en cabinet "
            "(200 K TND/an) et une rémunération de gérance de la clinique (150 K TND/an). "
            "La clinique est soumise à la TVA au taux réduit pour les actes médicaux. "
            "Il envisage d'apporter son cabinet à la SARL et de se salarier."
        ),
        objective=(
            "Comparer le régime IRPP (revenus libéraux vs rémunération gérance), "
            "la déductibilité des rémunérations du gérant en matière d'IS, "
            "le régime TVA des honoraires médicaux, et l'intérêt de la restructuration envisagée."
        ),
        topics=("irpp", "is", "gérance", "rémunération", "tva", "médecin", "honoraires"),
        extra_topics=("retenue", "source", "dividendes", "libéral"),
    ),
    ConsultationScenario(
        idx=8,
        reference="CONS-2024-008",
        case_label="hr_consultant_mix",
        client_type="mixte",
        client_name="Mme Fatma Chaabane / RH Partners SARL",
        client_description=(
            "Mme Fatma Chaabane est consultante RH indépendante. Elle a créé RH Partners SARL "
            "en 2022. Elle facture ses missions via la société mais perçoit également des "
            "honoraires à titre personnel d'entreprises qui préfèrent contractualiser "
            "directement avec elle."
        ),
        situation=(
            "En 2024, RH Partners SARL a réalisé 320 K TND de CA (IS) et Mme Chaabane "
            "a perçu 90 K TND d'honoraires personnels (IRPP BNC). Elle est aussi "
            "administratrice d'une SA cotée et perçoit des jetons de présence. "
            "Les retenus à la source effectuées varient selon les clients payeurs."
        ),
        objective=(
            "Analyser la qualification fiscale des honoraires personnels vs revenus SARL, "
            "le régime des jetons de présence (retenue, IS, IRPP), "
            "les risques d'abus de droit pour double flux, "
            "et l'optimisation de la structure de rémunération."
        ),
        topics=("irpp", "is", "honoraires", "retenue", "jetons de présence", "bnc"),
        extra_topics=("abus de droit", "tva", "rémunération", "gérance"),
    ),
    ConsultationScenario(
        idx=9,
        reference="CONS-2024-009",
        case_label="digital_services_vat",
        client_type="entreprise",
        client_name="StreamTN (filiale tunisienne de StreamGlobal BV)",
        client_description=(
            "StreamGlobal BV est une plateforme néerlandaise de streaming vidéo. "
            "Elle a créé StreamTN, filiale tunisienne, pour facturer ses abonnements "
            "aux clients résidents en Tunisie depuis janvier 2024."
        ),
        situation=(
            "StreamTN facture des abonnements mensuels (15 TND/mois) à 45 000 abonnés "
            "tunisiens. Les services sont entièrement dématérialisés. La maison mère "
            "percevait auparavant ces abonnements directement depuis les Pays-Bas sans "
            "TVA tunisienne. La filiale s'interroge sur ses obligations TVA, IS et retenue."
        ),
        objective=(
            "Déterminer l'assujettissement TVA des services numériques (territorialité, "
            "taux applicable), le régime IS de la filiale, les obligations déclaratives "
            "et de facturation électronique, et la situation avant création de la filiale."
        ),
        topics=("tva", "services numériques", "territoire", "abonnement", "électronique"),
        extra_topics=("is", "retenue", "facturation", "non-résident", "prestation"),
    ),
    ConsultationScenario(
        idx=10,
        reference="CONS-2024-010",
        case_label="btp_etablissement_stable",
        client_type="entreprise",
        client_name="BuildCon Belgium SA (chantier Tunisie)",
        client_description=(
            "BuildCon Belgium SA est une société de BTP belge spécialisée dans la "
            "construction de centrales solaires. Elle a remporté un appel d'offres "
            "pour la construction d'une centrale photovoltaïque en Tunisie."
        ),
        situation=(
            "Les travaux ont débuté en mars 2024 et dureront 14 mois. BuildCon Belgium "
            "importe du matériel et fait intervenir 40 techniciens belges sur le chantier. "
            "Elle sous-traite 30 % des travaux à des entreprises tunisiennes locales. "
            "Elle n'a pas de structure juridique en Tunisie mais dispose d'une direction "
            "locale permanente sur le chantier."
        ),
        objective=(
            "Établir l'existence ou non d'un établissement stable en Tunisie "
            "(convention Belgique-Tunisie), le régime IS applicable, les retenues à la "
            "source sur paiements aux sous-traitants, les obligations TVA, douanières "
            "et les formalités administratives requises."
        ),
        topics=("établissement stable", "chantier", "convention", "belgique", "is", "retenue"),
        extra_topics=("tva", "importation", "sous-traitant", "btp", "travaux"),
    ),
)


# ---------------------------------------------------------------------------
# Neo4j context retrieval
# ---------------------------------------------------------------------------

class Neo4jContextService:
    """Retrieves legal text chunks from the Neo4j graph database."""

    _CHUNK_QUERY = """
    MATCH (c:Chunk)
    WHERE any(t IN $topics
              WHERE toLower(c.text) CONTAINS t
                 OR toLower(coalesce(c.article_ref,'')) CONTAINS t)
    WITH c,
         reduce(s = 0.0, t IN $topics |
             s + CASE WHEN toLower(c.text) CONTAINS t THEN 1.0 ELSE 0.0 END
         ) AS score
    RETURN c.doc_name     AS doc_name,
           coalesce(c.page_num, 0)        AS page_num,
           coalesce(c.article_ref, '')    AS article_ref,
           coalesce(c.section_title, '')  AS section_title,
           coalesce(c.text, '')           AS text,
           score
    ORDER BY score DESC, c.doc_name ASC, c.page_num ASC
    LIMIT $limit
    """

    def __init__(self) -> None:
        cfg = get_config()
        self.database = cfg["neo4j_database"]
        self.driver = None
        if GraphDatabase is not None:
            self.driver = GraphDatabase.driver(
                cfg["neo4j_uri"], auth=(cfg["neo4j_username"], cfg["neo4j_password"])
            )

    def close(self) -> None:
        if self.driver is not None:
            self.driver.close()

    def _run_chunk_query(self, topics: Sequence[str], limit: int) -> List[LegalSource]:
        if self.driver is None:
            return []
        with self.driver.session(database=self.database) as session:
            rows = session.run(
                self._CHUNK_QUERY,
                topics=[t.lower() for t in topics],
                limit=limit,
            )
            return [LegalSource(**dict(row)) for row in rows]

    def get_chunks_for_scenario(self, scenario: ConsultationScenario) -> List[LegalSource]:
        """Two-pass retrieval: primary topics (top 8) + secondary topics (top 6), deduplicated."""
        primary = self._run_chunk_query(scenario.topics, limit=8)
        secondary = self._run_chunk_query(
            scenario.topics + scenario.extra_topics, limit=6
        ) if scenario.extra_topics else []
        all_sources = primary + secondary
        seen: Dict[str, bool] = {}
        deduped: List[LegalSource] = []
        for s in all_sources:
            key = hashlib.sha256(s.text.strip().encode("utf-8")).hexdigest()[:_DEDUP_HASH_LENGTH]
            if key not in seen:
                seen[key] = True
                deduped.append(s)
        return deduped[:10]


# ---------------------------------------------------------------------------
# Azure OpenAI LLM service
# ---------------------------------------------------------------------------

class AzureOpenAIService:
    """Calls Azure OpenAI (same credentials as .NET GraphRagService)."""

    SYSTEM_PROMPT = (
        "Tu es un expert fiscal tunisien senior dans un grand cabinet de conseil international "
        "(EY, PwC, Deloitte). Tu rédiges des consultations fiscales professionnelles, "
        "précises et fondées sur les textes légaux tunisiens. "
        "Tu cites toujours les articles de loi, codes et conventions pertinents. "
        "Tu rédiges exclusivement en français, avec un niveau de langage juridique et fiscal élevé. "
        "Ne fabrique pas de références légales : utilise uniquement ce qui est présent dans les "
        "extraits fournis ou dans ta connaissance avérée du droit tunisien."
    )

    def __init__(self) -> None:
        self.model = os.getenv("LLM_MODEL", "gpt-4o")
        self.client: Optional[object] = None
        if AzureOpenAI is not None:
            api_key = os.getenv("OPENAI_API_KEY", "")
            endpoint = os.getenv("OPENAI_ENDPOINT", "")
            api_version = os.getenv("OPENAI_API_VERSION", "2024-02-15-preview")
            if api_key and endpoint:
                self.client = AzureOpenAI(
                    api_key=api_key,
                    azure_endpoint=endpoint,
                    api_version=api_version,
                )

    def is_available(self) -> bool:
        return self.client is not None

    def generate(self, user_prompt: str, temperature: float = 0.7) -> str:
        if self.client is None:
            raise RuntimeError(
                "Azure OpenAI client not available. "
                "Set OPENAI_API_KEY and OPENAI_ENDPOINT environment variables."
            )
        response = self.client.chat.completions.create(  # type: ignore[union-attr]
            model=self.model,
            messages=[
                {"role": "system", "content": self.SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            temperature=temperature,
        )
        return response.choices[0].message.content or ""


# ---------------------------------------------------------------------------
# GraphRAG context builder
# ---------------------------------------------------------------------------

def build_legal_context(sources: Sequence[LegalSource]) -> str:
    """Format Neo4j chunks into a readable context block for the LLM prompt."""
    if not sources:
        return "(Aucun extrait légal disponible depuis Neo4j — raisonner à partir des connaissances générales du droit tunisien.)"
    lines: List[str] = []
    for i, s in enumerate(sources, 1):
        ref = s.article_ref.strip() if s.article_ref else "—"
        title = s.section_title.strip() if s.section_title else ""
        lines.append(f"[Extrait {i} | {s.doc_name}, p.{s.page_num} | Art. {ref}{' | ' + title if title else ''}]")
        # Truncate very long chunks to keep prompt size reasonable
        text = s.text.strip()
        if len(text) > _MAX_CHUNK_TEXT_CHARS:
            text = text[:_MAX_CHUNK_TEXT_CHARS] + "…"
        lines.append(text)
        lines.append("")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# LLM prompt builder
# ---------------------------------------------------------------------------

SECTION_PROMPT_TEMPLATE = """\
=== DOSSIER DE CONSULTATION FISCALE ===

Référence          : {reference}
Client             : {client_name} ({client_type})
Description client : {client_description}

FAITS :
{situation}

OBJECTIF DE LA CONSULTATION :
{objective}

=== EXTRAITS LÉGAUX (GraphRAG Neo4j) ===
{legal_context}

=== INSTRUCTIONS ===
Rédige les 4 sections ci-dessous de la consultation fiscale. Chaque section doit être \
complète, professionnelle, et s'appuyer sur les extraits légaux fournis. \
Cite les articles de loi explicitement dans le texte (ex. « conformément à l'article 52 \
du CDPF », « en vertu de l'article 3 du CTVA »).

--- SECTION 1 : CONTEXTE ---
Rédige la section suivante :

1. CONTEXTE

1.1 Compréhension des faits
Nous comprenons que [développe les faits de manière narrative et professionnelle, \
présente le client, son activité, sa structure, les pays impliqués, \
les montants en jeu, les dates clés].

1.2 Étendue de nos travaux
Notre analyse portera sur les aspects suivants :
[liste de 6 à 10 points précis selon les enjeux du dossier — IS, IRPP, TVA, \
retenues à la source, prix de transfert, change, sécurité sociale, établissement stable, \
obligations déclaratives, etc. — avec références légales tirées des extraits]

--- SECTION 2 : SOMMAIRE EXÉCUTIF ---
Rédige la section suivante (c'est la section 3 de la consultation) :

3. Sommaire Exécutif
[3 à 5 paragraphes de synthèse des enjeux et conclusions principales, \
en citant les textes légaux. Chaque conclusion doit être fondée sur un extrait ou \
une règle de droit précise.]

Tableau de risques fiscaux :
| Risque identifié | Probabilité | Impact | Mesure recommandée |
|---|---|---|---|
[4 à 6 lignes avec risques spécifiques au dossier, probabilités (Faible/Moyenne/Élevée) \
et impacts (Faible/Moyen/Élevé/Critique)]

--- SECTION 3 : ANALYSES ---
Rédige la section suivante (c'est la section 4 de la consultation) :

4. Analyses détaillées
[Pour chaque enjeu identifié en 1.2, développe une analyse de 2 à 4 phrases \
citant les articles applicables. Utilise le format tableau Q&A ci-dessous.]

Tableau d'analyse :
| Question | Règle applicable (articles) | Analyse et conclusion |
|---|---|---|
[5 à 8 lignes couvrant les questions clés du dossier avec citations légales précises]

--- SECTION 4 : DOCUMENTS ET RÉFÉRENCES ---
Rédige la section suivante (c'est la section 5 de la consultation) :

5. Documents et Références

Sources légales utilisées :
[liste des documents, articles, conventions cités dans la consultation]

Abréviations :
[liste de toutes les abréviations utilisées avec leur développement complet]

=== FIN DES INSTRUCTIONS ===
"""


def build_prompt(scenario: ConsultationScenario, legal_context: str) -> str:
    return SECTION_PROMPT_TEMPLATE.format(
        reference=scenario.reference,
        client_name=scenario.client_name,
        client_type=scenario.client_type,
        client_description=scenario.client_description,
        situation=scenario.situation,
        objective=scenario.objective,
        legal_context=legal_context,
    )


# ---------------------------------------------------------------------------
# LLM response parser  → template tokens
# ---------------------------------------------------------------------------

_SECTION_MARKERS = {
    "CONTEXTE": [
        "--- SECTION 1 : CONTEXTE ---",
        "1. CONTEXTE",
        "1.CONTEXTE",
    ],
    "SOMMAIRE_EXECUTIF": [
        "--- SECTION 2 : SOMMAIRE EXÉCUTIF ---",
        "3. Sommaire Exécutif",
        "3. SOMMAIRE",
        "SOMMAIRE EXÉCUTIF",
    ],
    "ANALYSES": [
        "--- SECTION 3 : ANALYSES ---",
        "4. Analyses",
        "4. ANALYSES",
        "ANALYSES DÉTAILLÉES",
    ],
    "DOCUMENTS_REFERENCES": [
        "--- SECTION 4 : DOCUMENTS ET RÉFÉRENCES ---",
        "5. Documents et Références",
        "5. DOCUMENTS",
        "DOCUMENTS ET RÉFÉRENCES",
    ],
}

# Ordered list of token keys — used to split the response
_TOKEN_ORDER = ["CONTEXTE", "SOMMAIRE_EXECUTIF", "ANALYSES", "DOCUMENTS_REFERENCES"]


def _find_marker_pos(text: str, markers: List[str]) -> int:
    """Return the index of the first marker found (case-insensitive), or -1."""
    text_lower = text.lower()
    for m in markers:
        pos = text_lower.find(m.lower())
        if pos != -1:
            return pos
    return -1


def parse_llm_response(llm_text: str, scenario: ConsultationScenario) -> Dict[str, str]:
    """
    Split the LLM output into template token values.
    Falls back to the full text for CONTEXTE if markers are not found.
    """
    tokens: Dict[str, str] = {
        "REFERENCE": scenario.reference,
        "TYPE_CAS": scenario.case_label.replace("_", " ").title(),
    }

    # Find positions of each section marker
    positions: List[Tuple[str, int]] = []
    for key in _TOKEN_ORDER:
        pos = _find_marker_pos(llm_text, _SECTION_MARKERS[key])
        positions.append((key, pos))

    # Sort by position, filter out -1
    found = [(k, p) for k, p in positions if p != -1]
    found.sort(key=lambda x: x[1])

    if not found:
        # No markers found — put everything in CONTEXTE as fallback
        tokens["CONTEXTE"] = llm_text.strip()
        tokens["SOMMAIRE_EXECUTIF"] = ""
        tokens["ANALYSES"] = ""
        tokens["DOCUMENTS_REFERENCES"] = ""
        return tokens

    for i, (key, start_pos) in enumerate(found):
        # Content starts right after the marker line
        marker_end = llm_text.find("\n", start_pos)
        content_start = marker_end + 1 if marker_end != -1 else start_pos

        # Content ends at the start of the next found section
        if i + 1 < len(found):
            content_end = found[i + 1][1]
            tokens[key] = llm_text[content_start:content_end].strip()
        else:
            tokens[key] = llm_text[content_start:].strip()

    # Ensure all keys are present
    for key in _TOKEN_ORDER:
        if key not in tokens:
            tokens[key] = ""

    return tokens


# ---------------------------------------------------------------------------
# docx output helpers
# ---------------------------------------------------------------------------

def slugify_filename(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    return re.sub(r"[^a-zA-Z0-9_-]+", "_", normalized).strip("_")


def replace_tokens_in_docx(template_path: Path, output_path: Path, tokens: Dict[str, str]) -> None:
    """Fill a .docx template's {{TOKEN}} placeholders with generated content."""
    if Document is None:
        raise RuntimeError("python-docx is not installed. Install it with: pip install python-docx")
    doc = Document(str(template_path))
    token_map = {f"{{{{{k}}}}}": v for k, v in tokens.items()}

    def _replace_in_runs(runs: list, token: str, value: str) -> bool:
        replaced = False
        for run in runs:
            if token in run.text:
                run.text = run.text.replace(token, value)
                replaced = True
        return replaced

    for para in doc.paragraphs:
        for token, value in token_map.items():
            if token in para.text and not _replace_in_runs(para.runs, token, value):
                para.text = para.text.replace(token, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for token, value in token_map.items():
                    if token in cell.text:
                        for para in cell.paragraphs:
                            if token in para.text and not _replace_in_runs(para.runs, token, value):
                                para.text = para.text.replace(token, value)

    doc.save(str(output_path))


def create_docx_from_scratch(output_path: Path, tokens: Dict[str, str]) -> None:
    """Create a formatted .docx when no template is available."""
    if Document is None:
        raise RuntimeError("python-docx is not installed. Install it with: pip install python-docx")

    doc = Document()

    # Title
    title_para = doc.add_heading(
        f"CONSULTATION FISCALE — {tokens.get('TYPE_CAS', '').upper()}", level=1
    )
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # type: ignore[union-attr]

    ref_para = doc.add_paragraph(f"Référence : {tokens.get('REFERENCE', '')}")
    ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # type: ignore[union-attr]
    doc.add_paragraph("")

    ordered = [
        ("CONTEXTE", "1. Contexte"),
        ("SOMMAIRE_EXECUTIF", "3. Sommaire Exécutif"),
        ("ANALYSES", "4. Analyses"),
        ("DOCUMENTS_REFERENCES", "5. Documents et Références"),
    ]

    for token_key, section_title in ordered:
        content = tokens.get(token_key, "").strip()
        if not content:
            continue
        doc.add_heading(section_title, level=2)
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Main report generation loop
# ---------------------------------------------------------------------------

def generate_reports(
    count: int,
    output_dir: Path,
    template_path: Optional[Path],
) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)

    # Limit count to available distinct scenarios
    if count > len(SCENARIOS):
        print(f"⚠️  count={count} exceeds {len(SCENARIOS)} distinct scenarios — using {len(SCENARIOS)}")
        count = len(SCENARIOS)

    neo4j = Neo4jContextService()
    llm = AzureOpenAIService()

    if not llm.is_available():
        print("⚠️  Azure OpenAI not configured (OPENAI_API_KEY / OPENAI_ENDPOINT missing).")
        print("    Reports will use placeholder content only.")

    try:
        for scenario in SCENARIOS[:count]:
            print(f"\n📄 [{scenario.idx}/{count}] {scenario.reference} — {scenario.client_name}")

            # Step 1 — Retrieve Neo4j legal chunks
            try:
                sources = neo4j.get_chunks_for_scenario(scenario)
                print(f"   🔍 Neo4j: {len(sources)} legal chunks retrieved")
            except Exception as exc:
                print(f"   ⚠️  Neo4j unavailable: {exc}")
                sources = []

            # Step 2 — Build GraphRAG context string
            legal_context = build_legal_context(sources)

            # Step 3 — Call LLM with Neo4j-grounded prompt
            if llm.is_available():
                print("   🤖 Calling LLM (Azure OpenAI)…")
                try:
                    prompt = build_prompt(scenario, legal_context)
                    llm_text = llm.generate(prompt, temperature=0.7)
                    print("   ✅ LLM response received")
                except Exception as exc:
                    print(f"   ❌ LLM call failed: {exc}")
                    llm_text = _fallback_content(scenario, sources)
            else:
                llm_text = _fallback_content(scenario, sources)

            # Step 4 — Parse response into template tokens
            tokens = parse_llm_response(llm_text, scenario)

            # Step 5 — Write .docx
            file_stem = slugify_filename(f"Consultation_{scenario.reference}_{scenario.case_label}")
            output_path = output_dir / f"{file_stem}.docx"

            if template_path and template_path.exists():
                replace_tokens_in_docx(template_path, output_path, tokens)
                print(f"   💾 Saved (template): {output_path}")
            else:
                create_docx_from_scratch(output_path, tokens)
                if template_path and not template_path.exists():
                    print(f"   ⚠️  Template not found at '{template_path}' — created bare docx")
                print(f"   💾 Saved: {output_path}")

    finally:
        neo4j.close()

    print(f"\n✅ Done — {count} consultation(s) written to '{output_dir}'")


def _fallback_content(scenario: ConsultationScenario, sources: Sequence[LegalSource]) -> str:
    """
    Minimal structured content used when the LLM is not available.
    Clearly flags that this is a placeholder requiring LLM enrichment.
    """
    legal_refs = "\n".join(
        f"- {s.short_ref()}" for s in sources
    ) or "- Aucune source Neo4j disponible"

    return (
        "--- SECTION 1 : CONTEXTE ---\n"
        f"1. CONTEXTE\n\n"
        f"1.1 Compréhension des faits\n"
        f"Nous comprenons que {scenario.client_name} ({scenario.client_type}) nous consulte "
        f"pour : {scenario.objective}\n\n"
        f"Faits : {scenario.situation}\n\n"
        "1.2 Étendue de nos travaux\n"
        "[À compléter par le LLM — configurer OPENAI_API_KEY et OPENAI_ENDPOINT]\n\n"
        "--- SECTION 2 : SOMMAIRE EXÉCUTIF ---\n"
        "3. Sommaire Exécutif\n"
        "[À compléter par le LLM]\n\n"
        "--- SECTION 3 : ANALYSES ---\n"
        "4. Analyses\n"
        "[À compléter par le LLM]\n\n"
        "--- SECTION 4 : DOCUMENTS ET RÉFÉRENCES ---\n"
        "5. Documents et Références\n\n"
        "Sources Neo4j :\n"
        f"{legal_refs}\n"
    )


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate 10 Neo4j-grounded Tunisian tax consultation reports (.docx)"
    )
    parser.add_argument(
        "--count",
        type=int,
        default=10,
        help="Number of consultations to generate (max 10, default: 10)",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("generated_consultations"),
        help="Output directory (default: generated_consultations/)",
    )
    parser.add_argument(
        "--template-path",
        type=Path,
        default=Path("template_fr.docx"),
        help="Path to the Word template file (default: template_fr.docx)",
    )
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    generate_reports(
        count=args.count,
        output_dir=args.output_dir,
        template_path=args.template_path,
    )
